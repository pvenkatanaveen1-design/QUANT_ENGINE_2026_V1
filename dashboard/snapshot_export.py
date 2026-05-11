"""
Snapshot text builders for dashboard export (Word / PDF).

Reads the same Redis keys as the main dashboard views. No Streamlit dependency.
"""

from __future__ import annotations

import json
from datetime import datetime
from io import BytesIO
from typing import Any, Callable

from core.bus import get_value
from core.config import load_config
from core.system_registry import TRACKED_PHASES, TRACKED_SYSTEMS, get_phase_status, get_system_status
from core.system_mode import get_system_mode
from market.features.atr import ATR_SUPPORTED_SYMBOLS
from risk.cost_guard import SUPPORTED_SYMBOLS as COST_SYMBOLS


def _ts(raw: Any) -> str:
    if raw is None:
        return "—"
    try:
        t = float(raw)
        if t > 10_000_000_000:
            t /= 1000
        return datetime.fromtimestamp(t).strftime("%Y-%m-%d %H:%M:%S")
    except (TypeError, ValueError, OSError):
        return str(raw)


def _json_block(obj: Any) -> str:
    try:
        return json.dumps(obj, indent=2, default=str, ensure_ascii=False)
    except TypeError:
        return str(obj)


# ---- section bodies -------------------------------------------------


def _sec_system_mode() -> str:
    return f"Current mode: {get_system_mode()}\n(TEST / LIVE from .env via core.system_mode)"


def _sec_engine_registry() -> str:
    lines = ["Phases:"]
    for p in TRACKED_PHASES:
        row = get_phase_status(p)
        lines.append(f"  {p}: {row.get('status')}")
    lines.append("Subsystems:")
    for name in TRACKED_SYSTEMS:
        row = get_system_status(name)
        lines.append(
            f"  {name}: {row.get('status')} | error={row.get('error')} | last_update={_ts(row.get('last_update'))}"
        )
    return "\n".join(lines)


def _sec_active_config() -> str:
    fb = load_config()
    mode = get_value("config:system_mode")
    if mode is None:
        mode = fb.get("SYSTEM_MODE")
    symbols = get_value("config:symbols")
    if symbols is None:
        symbols = fb.get("MT5_SYMBOLS", [])
    lines = [
        f"SYSTEM_MODE: {mode}",
        f"MT5_SYMBOLS: {symbols}",
        f"Redis config:last_update: {get_value('config:last_update')}",
        f"ACCOUNT_SIZE: {get_value('config:account_size') or fb.get('ACCOUNT_SIZE')}",
        f"daily_dd_warning: {get_value('config:daily_dd_warning') or fb.get('DAILY_DD_WARNING')}",
        f"daily_dd_block: {get_value('config:daily_dd_block') or fb.get('DAILY_DD_BLOCK')}",
        f"max_dd_block: {get_value('config:max_dd_block') or fb.get('MAX_DD_BLOCK')}",
        f"default_risk_per_trade: {get_value('config:default_risk_per_trade') or fb.get('DEFAULT_RISK_PER_TRADE')}",
        f"prop_firm: {get_value('config:prop_firm') or fb.get('PROP_FIRM_NAME')}",
        f"redis: {get_value('config:redis_host') or fb.get('REDIS_HOST')}:{get_value('config:redis_port') or fb.get('REDIS_PORT')}",
    ]
    return "\n".join(lines)


def _sec_legacy_system_status() -> str:
    raw = get_value("system:status")
    if raw is None:
        return "(no system:status key)"
    return _json_block(raw)


def _sec_heartbeat() -> str:
    keys = [
        "heartbeat:status",
        "heartbeat:last_check",
        "heartbeat:mt5_healthy",
        "heartbeat:market_fresh",
        "heartbeat:clock_fresh",
        "heartbeat:overall",
        "pulse:status",
        "mt5:connection",
    ]
    lines = []
    for k in keys:
        lines.append(f"{k} = {get_value(k)}")
    lines.append(f"market:XAUUSD:timestamp = {get_value('market:XAUUSD:timestamp')}")
    lines.append(f"clock:status = {get_value('clock:status')} | clock:utc = {get_value('clock:utc')}")
    return "\n".join(lines)


def _sec_active_risk_env() -> str:
    cfg = load_config()
    lines = [
        f"PROP_FIRM_NAME: {cfg.get('PROP_FIRM_NAME')}",
        f"ACCOUNT_SIZE: {cfg.get('ACCOUNT_SIZE')}",
        f"DEFAULT_RISK_PER_TRADE: {cfg.get('DEFAULT_RISK_PER_TRADE')}",
        f"DAILY_DD_WARNING: {cfg.get('DAILY_DD_WARNING')}",
        f"DAILY_DD_BLOCK: {cfg.get('DAILY_DD_BLOCK')}",
        f"MAX_DD_BLOCK: {cfg.get('MAX_DD_BLOCK')}",
        f"MAX_DD_APPROACH: {cfg.get('MAX_DD_APPROACH')}",
    ]
    return "\n".join(lines)


def _sec_risk_shield() -> str:
    lines = [
        f"account:balance = {get_value('account:balance')}",
        f"account:equity = {get_value('account:equity')}",
        f"risk:daily_dd = {get_value('risk:daily_dd')}",
        f"risk:max_dd = {get_value('risk:max_dd')}",
        f"risk:shield = {get_value('risk:shield')}",
        f"risk:status = {get_value('risk:status')}",
        f"risk:block_trading = {get_value('risk:block_trading')}",
        f"risk:last_update (local) = {_ts(get_value('risk:last_update'))}",
    ]
    return "\n".join(lines)


def _sec_drawdown() -> str:
    lines = [
        f"dd:start_balance = {get_value('dd:start_balance')}",
        f"dd:peak_equity = {get_value('dd:peak_equity')}",
        f"dd:lowest_equity = {get_value('dd:lowest_equity')}",
        f"dd:current_daily_dd = {get_value('dd:current_daily_dd')}",
        f"dd:max_daily_dd = {get_value('dd:max_daily_dd')}",
        f"dd:last_reset = {get_value('dd:last_reset')}",
        f"dd:tracker_status = {get_value('dd:tracker_status')}",
    ]
    return "\n".join(lines)


def _sec_news_guard() -> str:
    lines = [
        f"news:event = {get_value('news:event')}",
        f"news:event_time = {get_value('news:event_time')}",
        f"news:blackout = {get_value('news:blackout')}",
        f"news:blackout_reason = {get_value('news:blackout_reason')}",
        f"news:minutes_remaining = {get_value('news:minutes_remaining')}",
        f"news:last_update = {_ts(get_value('news:last_update'))}",
        f"news:status = {get_value('news:status')}",
    ]
    return "\n".join(lines)


def _sec_cost_guard() -> str:
    lines = [f"cost:last_update = {_ts(get_value('cost:last_update'))}"]
    for sym in COST_SYMBOLS:
        lines.append(
            f"{sym}: spread={get_value(f'cost:{sym}:spread')} | pips={get_value(f'cost:{sym}:spread_pips')} | "
            f"status={get_value(f'cost:{sym}:status')} | block={get_value(f'cost:{sym}:block_trading')} | "
            f"reason={get_value(f'cost:{sym}:reason')}"
        )
    return "\n".join(lines)


def _sec_router() -> str:
    lines = [
        f"router:last_symbol = {get_value('router:last_symbol')}",
        f"router:last_decision = {get_value('router:last_decision')}",
        f"router:last_reason = {get_value('router:last_reason')}",
        f"router:last_update = {_ts(get_value('router:last_update'))}",
        f"router:status = {get_value('router:status')}",
    ]
    return "\n".join(lines)


def _sec_broker_bridge() -> str:
    lines = [
        f"execution:last_symbol = {get_value('execution:last_symbol')}",
        f"execution:last_side = {get_value('execution:last_side')}",
        f"execution:last_volume = {get_value('execution:last_volume')}",
        f"execution:last_status = {get_value('execution:last_status')}",
        f"execution:last_ticket = {get_value('execution:last_ticket')}",
        f"execution:last_reason = {get_value('execution:last_reason')}",
        f"execution:last_update = {_ts(get_value('execution:last_update'))}",
        f"execution:bridge_status = {get_value('execution:bridge_status')}",
    ]
    return "\n".join(lines)


def _sec_trade_journal() -> str:
    lines = [
        f"journal:total_trades = {get_value('journal:total_trades')}",
        f"journal:last_status = {get_value('journal:last_status')}",
        f"journal:last_update = {_ts(get_value('journal:last_update'))}",
        f"journal:last_trade = {_json_block(get_value('journal:last_trade'))}",
    ]
    try:
        from journal.trade_logger import fetch_recent_trades

        recent = fetch_recent_trades(25)
        lines.append("Recent trades (up to 25):")
        for r in recent:
            lines.append(f"  {r}")
    except Exception as exc:  # noqa: BLE001
        lines.append(f"(SQLite recent trades unavailable: {exc})")
    return "\n".join(lines)


def _sec_order_tracker() -> str:
    summary = get_value("positions:summary")
    lines = [
        f"orders:active_count = {get_value('orders:active_count')}",
        f"orders:last_symbol = {get_value('orders:last_symbol')}",
        f"orders:last_ticket = {get_value('orders:last_ticket')}",
        f"orders:last_status = {get_value('orders:last_status')}",
        f"orders:last_update = {_ts(get_value('orders:last_update'))}",
        f"positions:summary = {_json_block(summary)}",
    ]
    return "\n".join(lines)


def _sec_trade_manager() -> str:
    lines = [
        f"trade_manager:last_ticket = {get_value('trade_manager:last_ticket')}",
        f"trade_manager:last_symbol = {get_value('trade_manager:last_symbol')}",
        f"trade_manager:last_action = {get_value('trade_manager:last_action')}",
        f"trade_manager:last_reason = {get_value('trade_manager:last_reason')}",
        f"trade_manager:last_update = {_ts(get_value('trade_manager:last_update'))}",
        f"trade_manager:status = {get_value('trade_manager:status')}",
        f"trade_manager:events = {_json_block(get_value('trade_manager:events'))}",
    ]
    return "\n".join(lines)


def _sec_position_sizer() -> str:
    lines = [
        f"sizing:last_symbol = {get_value('sizing:last_symbol')}",
        f"sizing:last_balance = {get_value('sizing:last_balance')}",
        f"sizing:last_risk_percent = {get_value('sizing:last_risk_percent')}",
        f"sizing:last_sl_pips = {get_value('sizing:last_sl_pips')}",
        f"sizing:last_lot_size = {get_value('sizing:last_lot_size')}",
        f"sizing:last_update = {_ts(get_value('sizing:last_update'))}",
    ]
    return "\n".join(lines)


def _sec_atr() -> str:
    lines = [f"features:atr:last_update = {_ts(get_value('features:atr:last_update'))}"]
    for sym in ATR_SUPPORTED_SYMBOLS:
        lines.append(
            f"{sym}: atr={get_value(f'features:{sym}:atr')} | "
            f"vol={get_value(f'features:{sym}:volatility_state')} | "
            f"status={get_value(f'features:{sym}:atr_status')}"
        )
    return "\n".join(lines)


def _sec_market_ticks() -> str:
    cfg = load_config()
    symbols = cfg.get("MT5_SYMBOLS", []) or []
    lines = []
    for symbol in sorted(set(str(s) for s in symbols)):
        bid = get_value(f"market:{symbol}:bid")
        ask = get_value(f"market:{symbol}:ask")
        spread = get_value(f"market:{symbol}:spread")
        ts = _ts(get_value(f"market:{symbol}:timestamp"))
        lines.append(f"{symbol}: bid={bid} ask={ask} spread={spread} ts={ts}")
    return "\n".join(lines) if lines else "(no MT5_SYMBOLS in config)"


def _sec_market_session() -> str:
    keys = [
        "clock:utc",
        "clock:ist",
        "clock:session",
        "clock:is_weekend",
        "clock:london",
        "clock:newyork",
        "clock:overlap",
        "clock:asia",
        "clock:status",
    ]
    return "\n".join(f"{k} = {get_value(k)}" for k in keys)


# id -> (title, builder)
EXPORT_SECTIONS: tuple[tuple[str, str, Callable[[], str]], ...] = (
    ("system_mode", "System mode", _sec_system_mode),
    ("engine_registry", "Engine & phase registry", _sec_engine_registry),
    ("active_config", "Active system configuration (Redis / .env)", _sec_active_config),
    ("legacy_system_status", "Legacy system:status JSON", _sec_legacy_system_status),
    ("heartbeat", "System health / heartbeat", _sec_heartbeat),
    ("active_risk_env", "Active risk configuration (.env)", _sec_active_risk_env),
    ("risk_shield", "Risk shield status", _sec_risk_shield),
    ("drawdown", "Drawdown tracker", _sec_drawdown),
    ("news_guard", "News guard", _sec_news_guard),
    ("cost_guard", "Cost guard", _sec_cost_guard),
    ("router", "Execution router", _sec_router),
    ("broker_bridge", "Broker bridge", _sec_broker_bridge),
    ("trade_journal", "Trade journal (Redis + SQLite)", _sec_trade_journal),
    ("order_tracker", "Live order tracker", _sec_order_tracker),
    ("trade_manager", "Trade manager", _sec_trade_manager),
    ("position_sizer", "Position sizer", _sec_position_sizer),
    ("atr_features", "ATR feature status", _sec_atr),
    ("market_ticks", "Live market ticks", _sec_market_ticks),
    ("market_session", "Market session (clock)", _sec_market_session),
)

_SECTION_BY_ID: dict[str, tuple[str, Callable[[], str]]] = {
    sid: (title, fn) for sid, title, fn in EXPORT_SECTIONS
}

LABEL_BY_ID: dict[str, str] = {sid: title for sid, title, _ in EXPORT_SECTIONS}
ID_BY_LABEL: dict[str, str] = {title: sid for sid, title, _ in EXPORT_SECTIONS}


def collect_selected_sections(selected_ids: list[str] | None) -> list[tuple[str, str]]:
    """
    Build (section title, body text) pairs in dashboard order.

    ``selected_ids`` None or empty list means **all** sections (full report).
    """
    if not selected_ids:
        order = [sid for sid, _, _ in EXPORT_SECTIONS]
    else:
        sel = set(selected_ids)
        order = [sid for sid, _, _ in EXPORT_SECTIONS if sid in sel]
    out: list[tuple[str, str]] = []
    for sid in order:
        title, fn = _SECTION_BY_ID[sid]
        try:
            body = fn().strip()
        except Exception as exc:  # noqa: BLE001
            body = f"(Error building section: {exc})"
        out.append((title, body))
    return out


def build_docx_bytes(sections: list[tuple[str, str]]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("QUANT ENGINE 2026 — dashboard export", level=0)
    doc.add_paragraph(f"Generated (local): {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Sections included: {len(sections)}")
    for title, body in sections:
        doc.add_heading(title, level=1)
        for line in body.splitlines():
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _pdf_safe_line(line: str) -> str:
    return line.encode("latin-1", "replace").decode("latin-1")


def build_pdf_bytes(sections: list[tuple[str, str]]) -> bytes:
    from fpdf import FPDF
    try:
        from fpdf.enums import XPos, YPos
    except ImportError:  # pragma: no cover
        XPos = YPos = None  # type: ignore[misc, assignment]

    class ExportPDF(FPDF):
        pass

    pdf = ExportPDF()
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()
    # fpdf2 needs explicit width; w=0 is not always "full width" across versions.
    w = float(getattr(pdf, "epw", pdf.w - pdf.l_margin - pdf.r_margin))

    kwargs: dict = {}
    if XPos is not None:
        kwargs["new_x"] = XPos.LMARGIN
        kwargs["new_y"] = YPos.NEXT

    def _mcell(h: float, text: str) -> None:
        pdf.multi_cell(w=w, h=h, txt=_pdf_safe_line(text), **kwargs)

    pdf.set_font("Helvetica", "B", 14)
    _mcell(8, "QUANT ENGINE 2026 — dashboard export")
    pdf.set_font("Helvetica", size=9)
    _mcell(5, datetime.now().strftime("Generated: %Y-%m-%d %H:%M:%S"))
    _mcell(5, f"Sections: {len(sections)}")
    pdf.ln(2)

    for title, body in sections:
        pdf.set_font("Helvetica", "B", 11)
        _mcell(6, title)
        pdf.set_font("Helvetica", size=8)
        for raw_line in body.splitlines():
            _mcell(4, raw_line)
        pdf.ln(1)

    return bytes(pdf.output())


def default_multiselect_labels() -> list[str]:
    return [title for _, title, _ in EXPORT_SECTIONS]
